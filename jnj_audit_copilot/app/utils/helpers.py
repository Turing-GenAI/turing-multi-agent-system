import glob
import json
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from langchain.schema import HumanMessage, SystemMessage
from retry import retry

from app.common.config import ACTIVITY_LIST_FILE, INPUT_DIR, SITE_DATA_INPUT_FILE_NAMES
from app.common.constants import (
    FINAL_OUTPUT_DOCX_FILENAME,
    FINAL_OUTPUT_PAGE_TITLE,
    FINDINGS_OUTPUT_FOLDER,
    INTERMEDIATE_PROCESSED_FILE_PATH,
    INTERMEDIATE_SUMMARY_DOCS_PATH,
    SGR_INTERMEDIATE_FOLDER,
    SGR_OUTPUT_FOLDER,
    site_id,
    trial_id,
    project_root,
)
from app.utils.log_setup import get_logger

# Get the same logger instance set up earlier
logger = get_logger()

document_persist = "document_persist"
processed_files_location = "processed_files"
summary_docs_location = "summary_docs"
warn_user_for_missing_site_areas_msg = ""


def generate_unique_activity_id(run_id, site_id, trial_id, activity_count, trigger_site_area):
    unique_activity_id = "_".join([trigger_site_area, "{:03d}".format(activity_count)])
    return "<activity_id#" + unique_activity_id + "> ###  "


class checkResources:
    def __init__(self):
        pass

    def check_file_exists(self, file_path):
        """
        Check if a specific .xlsx file exists.

        Parameters:
        - file_path (str): Path to the .xlsx file.

        Returns:
        - bool: True if file exists and is an .xlsx, otherwise False.
        """
        return os.path.isfile(file_path) and file_path.endswith(".xlsx")

    def check_xlsx_files_in_directory(self, directory_path):
        """
        Check if a directory contains any .xlsx files.

        Parameters:
        - directory_path (str): Path to the directory.

        Returns:
        - bool: True if the directory exists and contains .xlsx files, otherwise False.
        """
        if os.path.isdir(directory_path):
            xlsx_files = glob.glob(os.path.join(directory_path, "*.xlsx"))
            return bool(xlsx_files)
        return False

    def check_directory_exists(self, directory_path):
        """
        Check if a directory exists.

        Parameters:
        - directory_path (str): Path to the directory.

        Returns:
        - bool: True if the directory exists, otherwise False.
        """
        return os.path.isdir(directory_path)


# Function to combine text files and save to a .docx file with titles and
# centered headers
def combine_txt_files_to_docx(folder_path, run_id, output_filename, page_title):
    activity_findings_path = os.path.join(folder_path, run_id)
    os.makedirs(activity_findings_path, exist_ok=True)
    # Create a new Document object
    doc = Document()
    # Add the document title "Final Inspection Output" at the top
    doc_title = doc.add_paragraph()
    doc_title_run = doc_title.add_run(page_title)
    doc_title_run.bold = True
    doc_title_run.font.size = Pt(16)  # Set font size for the document title
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the document title

    # Add some spacing after the document title
    doc.add_paragraph()
    # Iterate over all .txt files in the folder
    for filename in sorted(os.listdir(activity_findings_path)):
        if filename.endswith(".txt") and not filename.startswith("conclusion"):
            file_path = os.path.join(activity_findings_path, filename)

            # Add title (filename) as a centered paragraph
            title = doc.add_paragraph()
            title_run = title.add_run(f"{filename.replace('.txt', '')}")  # Use the filename as the title
            title_run.bold = True
            title_run.font.size = Pt(14)  # Set font size for the title
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the title

            # Open and read each .txt file
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()
                # Add content of each .txt file to the document
                doc.add_paragraph(content)
                # Add a page break after the file content
                doc.add_page_break()

    # Save the combined content as a .docx file
    final_output_filepath = os.path.join(activity_findings_path, output_filename)
    doc.save(final_output_filepath)


def save_final_output():
    combine_txt_files_to_docx(
        FINDINGS_OUTPUT_FOLDER,
        FINAL_OUTPUT_DOCX_FILENAME,
        FINAL_OUTPUT_PAGE_TITLE,
    )


def load_site_inspection_data(json_path):
    with open(json_path, "r") as file:
        return json.load(file)


def create_directory_structure(run_id: str):
    # Define the folders to create under inputs/
    input_folders = [
        INTERMEDIATE_PROCESSED_FILE_PATH,
        SGR_INTERMEDIATE_FOLDER,
        INTERMEDIATE_SUMMARY_DOCS_PATH,
    ]

    # Define the folders to create under outputs/
    output_folders = [
        f"{FINDINGS_OUTPUT_FOLDER}/{run_id}",
        SGR_OUTPUT_FOLDER,
    ]

    # Create input directories
    for folder in input_folders:
        os.makedirs(folder, exist_ok=True)

    # Create output directories
    for folder in output_folders:
        os.makedirs(folder, exist_ok=True)


def read_file(file_path, file_format, index_col=None, sheet_name=None):
    """
    Reads a file and returns a DataFrame based on the specified format (csv or xlsx).

    Parameters:
        file_path (str): The path to the file to be read.
        file_format (str): The format of the file ('csv' or 'xlsx').
        index_col (int or str, optional): The column to use as the row labels of the DataFrame. Default is None.
        sheet_name (str or int, optional): The sheet name or index to read (for xlsx files). Default is None.

    Returns:
        pd.DataFrame: The loaded DataFrame if successful.
        None: If an error occurs during reading.
    """
    try:
        if file_format.lower() == "csv":
            df = pd.read_csv(file_path)
        elif file_format.lower() == "xlsx":
            # Constructing the arguments based on provided parameters
            if index_col is not None and sheet_name is not None:
                # Case with both index_col and sheet_name provided
                df = pd.read_excel(file_path, index_col=index_col, sheet_name=sheet_name)
            elif index_col is not None:
                # Case with only index_col provided
                df = pd.read_excel(file_path, index_col=index_col)
            elif sheet_name is not None:
                # Case with only sheet_name provided
                df = pd.read_excel(file_path, sheet_name=sheet_name)
            else:
                # Case with neither index_col nor sheet_name
                df = pd.read_excel(file_path)
        else:
            logger.error("Error: Unsupported file format. Please use 'csv' or 'xlsx'.")
            return None
        return df
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}, while reading file: {file_path}")
        return None


@retry(tries=2, delay=5)
def choose_relavant_columns(
    selfrag_prompts,
    model_with_required_column_structure,
    file_summary,
    question,
):
    template = selfrag_prompts["RELEVANT_COLUMNS_SELECTION_PROMPT"]
    messages = [
        SystemMessage(content="You are an expert in filtering data."),
        HumanMessage(content=template.format(file_summary=file_summary, question=question)),
    ]

    try:
        response = model_with_required_column_structure.invoke(messages)
        try:
            return response.columns
        except Exception as e:
            print(f'Fetched data from response using ["columns"], due to :{e}')
            return response["columns"]
    except Exception as e:
        logger.error(f"Could not fetch any columns due to {e}")
        return []


def clean_context(data, required_columns):
    # Split the data into rows
    rows = data.strip().split("\n\n\n")

    # First row is the headers
    headers = rows[0].strip().split("\n")
    headers = [header.strip().replace("\n", "") for header in headers]
    required_columns = [col.strip().replace("\n", "") for col in required_columns]

    # Remaining rows are the data
    data_rows = [row.strip().split("\n") for row in rows[1:]]

    # Convert data rows to a list of dictionaries
    data_dicts = [dict(zip(headers, row)) for row in data_rows]

    # 'Number_Days_Outstanding' SHOULD BE EMPTY WHEN 'End_date' is missing
    new_data_dicts = []
    for d in data_dicts:
        if d.get("End_Date") == "":
            d["Number_Days_Outstanding"] = ""
        new_data_dicts.append(d)
    data_dicts = new_data_dicts.copy()

    # keep only required columns
    data_dicts = [{k: v for k, v in d.items() if k in required_columns} for d in data_dicts]
    return data_dicts


def format_qa_pair(question, answer):
    """Format Q and A pair"""

    formatted_string = ""
    formatted_string += f"Question: {question}\nAnswer: {answer}\n\n"
    return formatted_string.strip()


def clean_graph_inputs(inputs):
    """
    Clean the graph inputs by removing any site areas that do not have corresponding activities in the activities list.

    Args:
        inputs (dict): The graph inputs to clean.

    Returns:
        dict: The cleaned graph inputs.
    """
    global warn_user_for_missing_site_areas_msg
    with open(ACTIVITY_LIST_FILE, "r") as f:
        site_areas = list(json.load(f).keys())
    unavailable_site_areas = []
    available_site_areas = []
    input_site_areas = inputs["trigger_list"][0]["site_areas"]
    for site_area in input_site_areas:
        if site_area != "SGR":
            if site_area not in site_areas:
                unavailable_site_areas.append(site_area)
            else:
                available_site_areas.append(site_area)
    if "SGR" in input_site_areas:
        available_site_areas = ["SGR"] + available_site_areas
    else:
        unavailable_site_areas.append("SGR")
    if len(unavailable_site_areas) > 0:
        warn_user_for_missing_site_areas_msg = (
            f"Warning: The following site areas ({', '.join(unavailable_site_areas)}) lack activities in 'activities.json'"
            "and have been excluded from the copilot run. "
            "Please add the necessary activities to include these site areas in future runs. "
            "The graph input is being updated to remove the unavailable areas."
        )
    inputs["trigger_list"][0]["site_areas"] = available_site_areas
    return inputs


def generate_input_filepaths_dict(
    INPUT_DIR,
    SITE_DATA_INPUT_FILE_NAMES,
    INTERMEDIATE_PROCESSED_FILE_PATH,
    INTERMEDIATE_SUMMARY_DOCS_PATH,
    site_id,
    trial_id,
):
    """
    Generates a dictionary with file paths for each site area based on the provided data.

    Args:
        INPUT_DIR (str): The base directory path for input files.
        SITE_DATA_INPUT_FILE_NAMES (dict): Dictionary with relative filenames for each site area.
        INTERMEDIATE_PROCESSED_FILE_PATH (str): The base directory path for intermediate processed files.
        INTERMEDIATE_SUMMARY_DOCS_PATH (str): The base directory path for intermediate summary documents.
        site_id (str): The ID of the site being processed.
        trial_id (str): The ID of the trial being processed.

    Returns:
        dict: Dictionary containing structured paths for each site area.
    """
    input_filepaths_dict = {}

    for site_area, data_filenames in SITE_DATA_INPUT_FILE_NAMES.items():
        # filenames for each site
        site_data_filename = data_filenames.get("site_data_filename", "")
        guidelines_data_filename = data_filenames.get("guidelines_data_filename", "")

        # Construct full file paths
        input_file_path = os.path.join(INPUT_DIR, site_data_filename)
        guidelines_pdf_path = os.path.join(INPUT_DIR, guidelines_data_filename)

        # Define processed directory path
        processed_dir = os.path.join(
            INTERMEDIATE_PROCESSED_FILE_PATH,
            f"{site_id}_{trial_id}",
            site_area,
        )

        # Paths for data and filtered directories
        filtered_data_dir = os.path.join(processed_dir, "data")
        filtered_filename = os.path.splitext(os.path.basename(site_data_filename))[0]

        # Construct paths and store in the dictionary
        input_filepaths_dict[site_area] = {
            "input_file_path": input_file_path,
            "filtered_input_folderpath": filtered_data_dir,
            "filtered_root_dir_path": os.path.join(
                filtered_data_dir,
                f"filtered_{filtered_filename}",
            ),
            "summary_docs_folder": os.path.join(
                INTERMEDIATE_SUMMARY_DOCS_PATH,
                f"{site_id}_{trial_id}",
                site_area,
            ),
            "summary_df_file_path": os.path.join(
                INTERMEDIATE_SUMMARY_DOCS_PATH,
                f"{site_id}_{trial_id}",
                site_area,
                f"filtered_{filtered_filename}_summary_df.xlsx",
            ),
            "guidelines_pdf_path": guidelines_pdf_path,
        }

    return input_filepaths_dict


# Generate filepaths dictionaries for ingestion and input
input_filepaths_dict = generate_input_filepaths_dict(
    INPUT_DIR,
    SITE_DATA_INPUT_FILE_NAMES,
    INTERMEDIATE_PROCESSED_FILE_PATH,
    INTERMEDIATE_SUMMARY_DOCS_PATH,
    site_id,
    trial_id,
)

def create_ingestion_filepaths_dict(
    activity_list_file,
    chromadb_dir,
    site_id,
    trial_id,
    summary_folder,
    document_folder,
    guidelines_folder,
):
    """
    Generates a dictionary with ingestion file paths for each site area.

    Args:
        activity_list_file (str): Path to the JSON file containing the list of site areas.
        chromadb_dir (str): The base directory path for the ChromaDB.
        site_id (str): The ID of the site.
        trial_id (str): The ID of the trial.
        summary_folder (str): The folder name for storing summary data.
        document_folder (str): The folder name for storing document data.
        guidelines_folder (str): The folder name for storing guidelines data.

    Returns:
        dict: Dictionary containing structured ingestion paths for each site area.
    """
    ingestion_filepaths_dict = {}

    with open(ACTIVITY_LIST_FILE, "r") as f:
        site_areas = list(json.load(f).keys())

    for area in site_areas:
        # Define reusable base directory for the current area
        chromadb_area_dir = os.path.join(chromadb_dir, f"{site_id}_{trial_id}", area)

        # Construct paths and store in the dictionary
        ingestion_filepaths_dict[area] = {
            "summary_persist_directory": os.path.join(chromadb_area_dir, summary_folder),
            "document_persist_directory": os.path.join(chromadb_area_dir, document_folder),
            "guidelines_persist_directory": os.path.join(chromadb_area_dir, guidelines_folder),
        }

    return ingestion_filepaths_dict

def create_ingestion_filepaths_dict_new(
    chromadb_dir,
    summary_folder,
    guidelines_folder,
    ):
    """
    Generates a dictionary with ingestion file paths for each site area.

    Args:
        chromadb_dir (str): The base directory path for the ChromaDB.
        summary_folder (str): The folder name for storing summary data.
        guidelines_folder (str): The folder name for storing guidelines data.

    Returns:
        dict: Dictionary containing structured ingestion paths for each site area.
    """
    ingestion_filepaths_dict = {}

    with open(ACTIVITY_LIST_FILE, "r") as f:
        site_areas = list(json.load(f).keys())

    for area in site_areas:
        # Define reusable base directory for the current area
        chromadb_area_dir = os.path.join(project_root, chromadb_dir, area)

        # Construct paths and store in the dictionary
        ingestion_filepaths_dict[area] = {
            "summary_persist_directory": os.path.join(chromadb_area_dir, summary_folder),
            "guidelines_persist_directory": os.path.join(chromadb_area_dir, guidelines_folder)
        }

    return ingestion_filepaths_dict

def is_folder_empty(folder_path):
    """
    Check if a folder is empty.
    
    Args:
        folder_path (str): Path to the folder
        
    Returns:
        bool: True if the folder is empty or doesn't exist, False otherwise
    """
    try:
        return len(os.listdir(folder_path)) == 0
    except FileNotFoundError:
        # If the folder doesn't exist, consider it empty
        return True